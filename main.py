import asyncio
import json
import os
import re
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple

import httpx
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt

from aiogram import Bot, Dispatcher, F, Router
from aiogram.client.default import DefaultBotProperties
from aiogram.enums import ParseMode
from aiogram.filters import CommandStart
from aiogram.exceptions import TelegramBadRequest
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import (
    BufferedInputFile,
    CallbackQuery,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
    KeyboardButton,
    Message,
    ReplyKeyboardMarkup,
    ReplyKeyboardRemove,
)

from db import DB
from ydisk import YDisk, sanitize_name

load_dotenv()

BOT_TOKEN = os.getenv("BOT_TOKEN")
ADMIN_IDS = {int(x.strip()) for x in (os.getenv("ADMIN_IDS") or "").split(",") if x.strip().isdigit()}
DB_PATH = os.getenv("DB_PATH", "storage/bot.db")

YANDEX_TOKEN = os.getenv("YANDEX_TOKEN")
YANDEX_ROOT = (os.getenv("YANDEX_ROOT") or "Фестиваль").strip().strip("/")
YANDEX_INBOX = (os.getenv("YANDEX_INBOX") or "INBOX").strip().strip("/")
YANDEX_LOCAL = (os.getenv("YANDEX_LOCAL") or "Локальные файлы").strip().strip("/")

if not BOT_TOKEN:
    raise RuntimeError("No BOT_TOKEN in .env")
if not ADMIN_IDS:
    raise RuntimeError("No ADMIN_IDS in .env")
if not YANDEX_TOKEN:
    raise RuntimeError("No YANDEX_TOKEN in .env")

db = DB(DB_PATH)
yd = YDisk(YANDEX_TOKEN)
router = Router()

COMMON_DL_MAP: dict[str, str] = {}
SUB_DL_MAP: dict[str, str] = {}

INTRO = (
    "Вас приветствует бот осветительного отдела Мастерской «12».\n"
    "Ниже расположено меню, в котором Вы можете посмотреть техническую информацию нашего отдела "
    "в разделе «Документы» и пройти опрос, который поможет нам подготовиться к Вашему приезду. "
    "Большая просьба пройти опрос!"
)

THANKS = (
    "Спасибо за уделенное время, Ваши ответы помогут нам лучше подготовиться к Вашему мероприятию!\n"
    "Вы можете изменить ответы, используя меню ниже.\n"
    "Так же Вы можете прислать фото и документы - райдер, лайтплот и т.д. Они будут храниться в меню «Документы»."
)

DATE_RE = re.compile(r"^\d{2}\.\d{2}\.\d{4}$")
PHONE_RE = re.compile(r"^[\d\+\-\(\) ]{6,}$")


def is_admin(uid: int) -> bool:
    return uid in ADMIN_IDS


def norm_date(s: str) -> Optional[str]:
    s = (s or "").strip()
    if not DATE_RE.match(s):
        return None
    try:
        datetime.strptime(s, "%d.%m.%Y")
        return s
    except ValueError:
        return None


def int_pos(s: str) -> Optional[int]:
    s = (s or "").strip()
    if not s.isdigit():
        return None
    n = int(s)
    return n if n > 0 else None


def norm_phone(s: str) -> Optional[str]:
    s = (s or "").strip()
    if not s:
        return None
    if not PHONE_RE.match(s):
        return None
    return re.sub(r"\s+", " ", s)


def folder_for(event_date: str, org: str, event_title: str) -> str:
    return f"{YANDEX_ROOT}/{sanitize_name(event_date)}-{sanitize_name(org)}-{sanitize_name(event_title)}"


def inbox_for(uid: int) -> str:
    return f"{YANDEX_ROOT}/{YANDEX_INBOX}/{uid}"


SURVEY_OPTIONS: Dict[str, List[str]] = {
    "scene": ["Большой зал", "Малый зал"],
    "night_mount": ["Да", "Нет"],
    "mount_who": ["Сами", "Сотрудники Мастерской «12»", "Совместно"],
    "extra_equipment": ["Нет", "Привезем свое", "Возьмем в прокате"],
    "power_type": ["Нет", "63А - 5 Pin", "32A - 5 Pin", "32A - 3 Pin"],
    "power_where": ["Левая сторона", "Арьер", "Правая сторона", "Авансцена"],
    "dimmer_needed": ["Да", "Нет"],
    "sfx": ["Нет", "Дым/туман", "Мыльные пузыри", "Снег", "Пиротехника", "Другое"],
    "operator": ["Сами", "Оператор Мастерской «12»"],
    "console_help": ["Да", "Нет", "Привезем свой пульт"],
    "confirm": ["✅ Сохранить", "🔁 Заново"],
}


def ikb(rows: List[List[Tuple[str, str]]]) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[[InlineKeyboardButton(text=t, callback_data=d) for (t, d) in row] for row in rows]
    )


def kb_inline(field: str, cols: int = 2) -> InlineKeyboardMarkup:
    opts = SURVEY_OPTIONS[field]
    rows: List[List[Tuple[str, str]]] = []
    row: List[Tuple[str, str]] = []
    for i, opt in enumerate(opts):
        row.append((opt, f"ans:{field}:{i}"))
        if len(row) >= cols:
            rows.append(row)
            row = []
    if row:
        rows.append(row)
    return ikb(rows)


def kb_power_types_multi(selected: List[str], none_selected: bool = False) -> InlineKeyboardMarkup:
    opts = SURVEY_OPTIONS["power_type"]
    rows: List[List[InlineKeyboardButton]] = []
    # ❌ показываем только если пользователь явно выбрал «Нет»
    text_none = "❌ Нет" if (none_selected and not selected) else "Нет"
    rows.append([InlineKeyboardButton(text=text_none, callback_data="pt:none")])
    for i, opt in enumerate(opts):
        if opt == "Нет":
            continue
        mark = "✅ " if opt in selected else "☑️ "
        rows.append([InlineKeyboardButton(text=f"{mark}{opt}", callback_data=f"pt:opt:{i}")])
    rows.append([InlineKeyboardButton(text="➡️ Далее", callback_data="pt:done")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def kb_sfx_multi(selected: List[str], none_selected: bool = False) -> InlineKeyboardMarkup:
    opts = SURVEY_OPTIONS["sfx"]
    rows: List[List[InlineKeyboardButton]] = []
    # ❌ показываем только если пользователь явно выбрал «Нет»
    text_none = "❌ Нет" if (none_selected and not selected) else "Нет"
    rows.append([InlineKeyboardButton(text=text_none, callback_data="sx:none")])
    for i, opt in enumerate(opts):
        if opt == "Нет":
            continue
        mark = "✅ " if opt in selected else "☑️ "
        rows.append([InlineKeyboardButton(text=f"{mark}{opt}", callback_data=f"sx:opt:{i}")])
    rows.append([InlineKeyboardButton(text="➡️ Далее", callback_data="sx:done")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def kb_survey_reply() -> ReplyKeyboardMarkup:
    return ReplyKeyboardMarkup(keyboard=[[KeyboardButton(text="⏸ Прервать и доделать позже")]], resize_keyboard=True)


def kb_menu(uid: int) -> ReplyKeyboardMarkup:
    rows = [
        [KeyboardButton(text="📝 Пройти опрос")],
        [KeyboardButton(text="📁 Документы"), KeyboardButton(text="📄 Мои ответы")],
        [KeyboardButton(text="✏️ Изменить ответы")],
    ]
    if db.get_draft(uid):
        rows.insert(1, [KeyboardButton(text="▶️ Продолжить опрос")])
    if is_admin(uid):
        rows.insert(0, [KeyboardButton(text="🛠 Админ")])
    return ReplyKeyboardMarkup(keyboard=rows, resize_keyboard=True)


def kb_admin_menu() -> ReplyKeyboardMarkup:
    rows = [
        [KeyboardButton(text="📋 Анкеты"), KeyboardButton(text="📊 Статистика")],
        [KeyboardButton(text="📄 Word анкеты"), KeyboardButton(text="🗑 Удалить анкету")],
        [KeyboardButton(text="⬅️ Назад")],
    ]
    return ReplyKeyboardMarkup(keyboard=rows, resize_keyboard=True)


def answers_text(a: Dict[str, Any]) -> str:
    def fmt_list(xs: List[str]) -> str:
        if not xs:
            return "—"
        return "\n" + "\n".join([f"  - {x}" for x in xs])

    def g(k: str) -> str:
        v = a.get(k)
        if v is None or v == "" or v == []:
            return "—"
        if isinstance(v, list):
            return fmt_list([str(x) for x in v])
        return str(v)

    power_where = a.get("power_where_list") or []
    power_type = str(a.get("power_type") or "").strip()
    power_items = a.get("power_items") or []
    power_types = a.get("power_types") or []

    if (not power_where) and power_items:
        flat: List[str] = []
        for it in power_items:
            t = (it or {}).get("type")
            for w in ((it or {}).get("where") or []):
                if t:
                    flat.append(f"{t}: {w}")
        power_where = flat

    if (not power_type) and power_items:
        power_type = ", ".join([str((it or {}).get("type")) for it in power_items if (it or {}).get("type")])

    if (not power_type) and power_types:
        power_type = ", ".join([str(x) for x in power_types if x])

    power_needed = "Да"
    if (not power_where) and (power_type in {"", "—", "Нет", "0"}):
        power_needed = "Нет"
    if str(power_type).strip() == "Нет":
        power_needed = "Нет"
    if power_needed == "Нет" and (power_types or power_items):
        power_needed = "Да"

    dimmer_needed = str(a.get("dimmer_needed") or "—").strip()

    sfx_list = a.get("sfx_list")
    if sfx_list is None:
        try:
            sfx_list = json.loads(a.get("sfx_json") or "[]")
        except Exception:
            sfx_list = []
    if not isinstance(sfx_list, list):
        sfx_list = []
    sfx_other = str(a.get("sfx_other") or "").strip()
    if sfx_other and ("Другое" in sfx_list):
        sfx_list = [x if x != "Другое" else f"Другое: {sfx_other}" for x in sfx_list]
    a["sfx_list"] = sfx_list

    operator = str(a.get("operator") or "—").strip()
    console_help = str(a.get("console_help") or "—").strip()
    scene = str(a.get("scene") or "").strip()

    if operator == "Оператор Мастерской «12»":
        console_model = "—"
    else:
        if console_help == "Привезем свой пульт":
            console_model = str(a.get("console_model") or "—").strip()
        else:
            console_model = "GrandMa2 Light" if scene == "Большой зал" else "Chamsys MQ500"

    lines = [
        f"1) Организация: {g('org')}",
        f"2) Должность: {g('role')}",
        f"3) Имя: {g('name')}",
        f"4) Телефон: {g('phone')}",
        f"5) Дата: {g('event_date')}",
        f"6) Название мероприятия: {g('event_title')}",
        f"7) Сцена: {g('scene')}",
        f"8) Ночной монтаж: {g('night_mount')}",
        f"9) Монтаж: {g('mount_who')}",
        f"10) Техников: {g('techs_count')}",
        f"11) Доп. оборудование: {g('extra_equipment')}",
        f"12) Вилки: {g('plugs')}",
        f"13) Силовые: {power_needed}",
        f"14) Где силовые: {fmt_list(power_where) if power_where else '—'}",
        f"15) Диммер: {dimmer_needed}",
    ]
    if dimmer_needed == "Да":
        lines.append(f"16) Диммер где и сколько: {g('dimmer_text')}")
    lines.append(f"17) Спецэффекты: {g('sfx_list')}")
    lines.extend(
        [
            f"18) Кто ведет: {operator}",
            f"19) Помощь с пультом: {console_help}",
            f"20) Пульт: {console_model}",
        ]
    )
    return "\n".join(lines)


EDIT_FIELDS = [
    ("org", "Организация"),
    ("role", "Должность"),
    ("name", "Имя"),
    ("phone", "Телефон"),
    ("event_date", "Дата (ДД.ММ.ГГГГ)"),
    ("event_title", "Название мероприятия"),
    ("scene", "Сцена"),
    ("night_mount", "Ночной монтаж"),
    ("mount_who", "Кто монтирует"),
    ("techs_count", "Сколько техников"),
    ("extra_equipment", "Доп. оборудование"),
    ("plugs", "Вилки"),
    ("power_block", "Силовые подключения"),
    ("dimmer_needed", "Диммер нужен"),
    ("dimmer_text", "Диммер: где и сколько"),
    ("operator", "Кто ведет"),
    ("console_help", "Помощь с пультом"),
    ("console_model", "Пульт модель"),
]

EDIT_OPTIONS: Dict[str, List[str]] = {
    "scene": SURVEY_OPTIONS["scene"],
    "night_mount": SURVEY_OPTIONS["night_mount"],
    "mount_who": SURVEY_OPTIONS["mount_who"],
    "extra_equipment": SURVEY_OPTIONS["extra_equipment"],
    "dimmer_needed": SURVEY_OPTIONS["dimmer_needed"],
    "operator": SURVEY_OPTIONS["operator"],
    "console_help": SURVEY_OPTIONS["console_help"],
}


def kb_edit_fields() -> ReplyKeyboardMarkup:
    kb_rows = [[KeyboardButton(text=title)] for _, title in EDIT_FIELDS]
    kb_rows.append([KeyboardButton(text="⬅️ Назад")])
    return ReplyKeyboardMarkup(keyboard=kb_rows, resize_keyboard=True)


def kb_reply_options(options: List[str]) -> ReplyKeyboardMarkup:
    kb_rows = [[KeyboardButton(text=o)] for o in options]
    kb_rows.append([KeyboardButton(text="⬅️ Назад")])
    return ReplyKeyboardMarkup(keyboard=kb_rows, resize_keyboard=True)


class Survey(StatesGroup):
    org = State()
    role = State()
    name = State()
    event_date = State()
    event_title = State()
    scene = State()
    night_mount = State()
    mount_who = State()
    techs_count = State()
    extra_equipment = State()
    plugs = State()
    power_type = State()
    power_count = State()
    power_where = State()
    dimmer_needed = State()
    dimmer_text = State()
    sfx = State()
    sfx_other = State()
    operator = State()
    console_help = State()
    console_model = State()
    phone = State()
    confirm = State()


class EditPower(StatesGroup):
    power_type = State()
    power_count = State()
    power_where = State()
    confirm = State()


class Edit(StatesGroup):
    pick = State()
    value = State()


class AdminWord(StatesGroup):
    pick_month = State()
    pick_form = State()


class AdminDel(StatesGroup):
    pick_month = State()
    pick_form = State()
    confirm = State()


class AdminForms(StatesGroup):
    pick_month = State()
    pick_form = State()


async def draft_get(state: FSMContext) -> Dict[str, Any]:
    d = await state.get_data()
    return d.get("draft") or {}


async def draft_set(state: FSMContext, patch: Dict[str, Any]) -> Dict[str, Any]:
    d = await draft_get(state)
    d.update(patch)
    await state.update_data(draft=d)
    return d


def _safe_row_get(row: Any, key: str, default: Any = None) -> Any:
    try:
        if hasattr(row, "keys") and key in row.keys():
            return row[key]
    except Exception:
        pass
    return default


def submission_to_dict(row: Any) -> Dict[str, Any]:
    power_where_list = json.loads(_safe_row_get(row, "power_where_json", "[]") or "[]")
    sfx_json = _safe_row_get(row, "sfx_json", "[]") or "[]"
    try:
        sfx_list = json.loads(sfx_json)
    except Exception:
        sfx_list = []
    if not isinstance(sfx_list, list):
        sfx_list = []
    sfx_other = _safe_row_get(row, "sfx_other", "") or ""
    return {
        "org": row["org"],
        "role": row["role"],
        "name": row["name"],
        "phone": _safe_row_get(row, "phone", "") or "",
        "event_date": row["event_date"],
        "event_title": row["event_title"],
        "scene": row["scene"],
        "night_mount": row["night_mount"],
        "mount_who": row["mount_who"],
        "techs_count": row["techs_count"],
        "extra_equipment": row["extra_equipment"],
        "plugs": row["plugs"],
        "power_type": (row["power_type"] or "Нет"),
        "power_where_list": power_where_list,
        "dimmer_needed": row["dimmer_needed"],
        "dimmer_text": row["dimmer_text"],
        "sfx_json": sfx_json,
        "sfx_other": sfx_other,
        "sfx_list": sfx_list,
        "operator": row["operator"],
        "console_help": row["console_help"],
        "console_model": row["console_model"],
        "ydisk_folder": row["ydisk_folder"],
        "id": row["id"],
        "user_id": row["user_id"],
    }


def build_docx_for_submission(sub: Any) -> bytes:
    d = Document()
    p = d.add_paragraph("Анкета участника фестиваля")
    if p.runs:
        p.runs[0].font.size = Pt(14)

    d.add_paragraph(f"ID анкеты: {sub['id']}")
    d.add_paragraph(f"Организация: {sub['org']}")
    d.add_paragraph(f"Название мероприятия: {sub['event_title']}")
    d.add_paragraph(f"Дата мероприятия: {sub['event_date']}")
    d.add_paragraph(f"Сцена: {sub['scene']}")
    d.add_paragraph("")

    table = d.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Вопрос"
    hdr[1].text = "Ответ"

    def add(q: str, a: str):
        row = table.add_row().cells
        row[0].text = q
        row[1].text = a if (a and str(a).strip()) else "—"

    add("1 Организация", sub["org"])
    add("2 Должность", sub["role"])
    add("3 Имя", sub["name"])
    add("4 Телефон", sub.get("phone", "") or "—")
    add("5 Дата проведения", sub["event_date"])
    add("6 Название мероприятия", sub["event_title"])
    add("7 Сцена", sub["scene"])
    add("8 Ночной монтаж", sub["night_mount"])
    add("9 Кто монтирует", sub["mount_who"])
    add("10 Сколько техников", sub["techs_count"])
    add("11 Доп. оборудование", sub["extra_equipment"])
    add("12 Вилки", sub["plugs"])

    power_where_list = json.loads(sub.get("power_where_json") or "[]")
    power_type = str(sub.get("power_type") or "").strip()
    power_needed = "Да"
    if (not power_where_list) and (power_type in {"", "—", "Нет", "0"}):
        power_needed = "Нет"
    if power_type == "Нет":
        power_needed = "Нет"
    add("13 Силовые подключения", power_needed)
    add("14 Где силовые", "\n".join(power_where_list) if power_where_list else "—")

    add("15 Диммер", sub["dimmer_needed"])
    if str(sub["dimmer_needed"]).strip() == "Да":
        add("16 Диммер где и сколько", sub["dimmer_text"])

    try:
        sfx_list = json.loads(sub.get("sfx_json") or "[]")
    except Exception:
        sfx_list = []
    if not isinstance(sfx_list, list):
        sfx_list = []
    sfx_other = str(sub.get("sfx_other") or "").strip()
    if sfx_other and ("Другое" in sfx_list):
        sfx_list = [x if x != "Другое" else f"Другое: {sfx_other}" for x in sfx_list]
    add("17 Спецэффекты", "\n".join(sfx_list) if sfx_list else "Нет")

    add("18 Кто ведет", sub["operator"])
    add("19 Помощь с пультом", sub["console_help"])
    add("20 Пульт", sub["console_model"])
    add("Папка на Я.Диске", sub["ydisk_folder"])

    d.add_paragraph("")
    d.add_paragraph("Подпись: ______________________")
    d.add_paragraph("Расшифровка: __________________")

    buf = BytesIO()
    d.save(buf)
    return buf.getvalue()


async def put_to_yandex(uid: int, file_name: str, data: bytes) -> str:
    sub_id, folder = db.get_user_last(uid)
    if not folder:
        folder = inbox_for(uid)
        await yd.ensure_folder(f"{YANDEX_ROOT}")
        await yd.ensure_folder(f"{YANDEX_ROOT}/{YANDEX_INBOX}")
        await yd.ensure_folder(folder)
    disk_path = f"{folder}/{sanitize_name(file_name)}"
    await yd.upload_bytes(disk_path, data, overwrite=True)
    db.save_doc(uid, sub_id, sanitize_name(file_name), disk_path)
    return disk_path


async def download_and_send(message: Message, ydisk_path: str):
    href = await yd.get_download_url(ydisk_path)
    if not href:
        await message.answer("Не удалось получить ссылку")
        return
    async with httpx.AsyncClient(timeout=120, follow_redirects=True) as client:
        r = await client.get(href)
        r.raise_for_status()
        data = r.content
    name = ydisk_path.split("/")[-1] or "file"
    await message.answer_document(BufferedInputFile(data, filename=name))


@router.message(CommandStart())
async def start(message: Message, state: FSMContext):
    await state.clear()
    await message.answer(INTRO, reply_markup=kb_menu(message.from_user.id))


@router.message(F.text == "📝 Пройти опрос")
async def m_survey(message: Message, state: FSMContext):
    if db.get_draft(message.from_user.id):
        kb = ReplyKeyboardMarkup(
            keyboard=[
                [KeyboardButton(text="▶️ Продолжить опрос"), KeyboardButton(text="🆕 Начать заново")],
                [KeyboardButton(text="⬅️ Назад")],
            ],
            resize_keyboard=True,
        )
        await state.clear()
        await message.answer("Есть незавершенный опрос. Что делаем?", reply_markup=kb)
        return
    await state.clear()
    await draft_set(
        state,
        {"power_types": [], "power_items": [], "power_i": 0, "power_none": False, "sfx_list": [], "sfx_other": "",
         "sfx_none": False},
    )
    await state.set_state(Survey.org)
    await message.answer("1) Как называется ваша организация?", reply_markup=kb_survey_reply())


@router.message(F.text == "🆕 Начать заново")
async def m_new(message: Message, state: FSMContext):
    db.delete_draft(message.from_user.id)
    await state.clear()
    await draft_set(
        state,
        {"power_types": [], "power_items": [], "power_i": 0, "power_none": False, "sfx_list": [], "sfx_other": "",
         "sfx_none": False},
    )
    await state.set_state(Survey.org)
    await message.answer("1) Как называется ваша организация?", reply_markup=kb_survey_reply())


@router.message(F.text == "▶️ Продолжить опрос")
async def m_resume(message: Message, state: FSMContext):
    row = db.get_draft(message.from_user.id)
    if not row:
        await message.answer("Черновика нет.", reply_markup=kb_menu(message.from_user.id))
        return

    await state.clear()
    try:
        d = json.loads(row["draft_json"])
    except Exception:
        d = {"power_types": [], "power_items": [], "power_i": 0, "sfx_list": [], "sfx_other": ""}
    d.setdefault("power_none", False)
    d.setdefault("sfx_none", False)
    await state.update_data(draft=d)
    await state.set_state(row["fsm_state"])
    st = row["fsm_state"]
    d2 = await draft_get(state)

    if st == Survey.org.state:
        return await message.answer("1) Как называется ваша организация?", reply_markup=kb_survey_reply())
    if st == Survey.role.state:
        return await message.answer("2) Ваша должность?", reply_markup=kb_survey_reply())
    if st == Survey.name.state:
        return await message.answer("3) Ваше имя?", reply_markup=kb_survey_reply())
    if st == Survey.event_date.state:
        return await message.answer("4) Дата проведения мероприятия? (ДД.ММ.ГГГГ)", reply_markup=kb_survey_reply())
    if st == Survey.event_title.state:
        return await message.answer("5) Название мероприятия?", reply_markup=kb_survey_reply())
    if st == Survey.scene.state:
        return await message.answer("6) На какой сцене будет проходить мероприятие?", reply_markup=kb_inline("scene", 2))
    if st == Survey.night_mount.state:
        return await message.answer("7) Нужен ли ночной монтаж перед мероприятием?", reply_markup=kb_inline("night_mount", 2))
    if st == Survey.mount_who.state:
        return await message.answer("8) Кто производит монтаж световой аппаратуры?", reply_markup=kb_inline("mount_who", 1))
    if st == Survey.techs_count.state:
        return await message.answer("9) Сколько техников на монтаж понадобится? (число)", reply_markup=kb_survey_reply())
    if st == Survey.extra_equipment.state:
        return await message.answer("10) Используете ли доп. световое оборудование?", reply_markup=kb_inline("extra_equipment", 1))
    if st == Survey.plugs.state:
        return await message.answer("11) Какие вилки на ваших приборах?", reply_markup=kb_survey_reply())
    if st == Survey.power_type.state:
        sel = d2.get("power_types") or []
        none_selected = bool(d2.get("power_none"))
        return await message.answer(
            "12) Выберите один или несколько вариантов и нажмите «➡️ Далее»:",
            reply_markup=kb_power_types_multi(sel, none_selected=none_selected),
        )
    if st == Survey.power_count.state:
        items = d2.get("power_items") or []
        i = int(d2.get("power_i") or 0)
        cur = items[i]["type"] if i < len(items) else ""
        return await message.answer(f"13) Сколько подключений нужно для «{cur}»?", reply_markup=kb_survey_reply())
    if st == Survey.power_where.state:
        items = d2.get("power_items") or []
        i = int(d2.get("power_i") or 0)
        if i >= len(items):
            await state.set_state(Survey.dimmer_needed)
            return await message.answer("15) Нужен ли диммер?", reply_markup=kb_inline("dimmer_needed", 2))
        cur = items[i]
        need = int(cur.get("count") or 0)
        got = len(cur.get("where") or [])
        left = max(need - got, 0)
        txt = f"14) Где нужны подключения для «{cur['type']}»?"
        if got > 0:
            txt += f" Осталось {left}"
        return await message.answer(txt, reply_markup=kb_inline("power_where", 2))
    if st == Survey.dimmer_needed.state:
        return await message.answer("15) Нужен ли диммер?", reply_markup=kb_inline("dimmer_needed", 2))
    if st == Survey.dimmer_text.state:
        return await message.answer("16) Диммер: где и сколько?", reply_markup=kb_survey_reply())
    if st == Survey.sfx.state:
        sel = d2.get("sfx_list") or []
        none_selected = bool(d2.get("sfx_none"))
        return await message.answer(
            "17) Выберите один или несколько вариантов и нажмите «➡️ Далее»:",
            reply_markup=kb_sfx_multi(sel, none_selected=none_selected),
        )
    if st == Survey.sfx_other.state:
        return await message.answer("Введите свой вариант спецэффектов:", reply_markup=kb_survey_reply())
    if st == Survey.operator.state:
        return await message.answer("18) Кто ведет мероприятие?", reply_markup=kb_inline("operator", 1))
    if st == Survey.console_help.state:
        console_name = "GrandMa2 Light" if d2.get("scene") == "Большой зал" else "Chamsys MQ500"
        return await message.answer(
            f"19) Мы используем пульт {console_name}. Нужна помощь с пультом?",
            reply_markup=kb_inline("console_help", 1),
        )
    if st == Survey.console_model.state:
        return await message.answer("20) Марка и модель вашего пульта? (текст)", reply_markup=kb_survey_reply())
    if st == Survey.phone.state:
        return await message.answer("21) Номер телефона для связи?", reply_markup=kb_survey_reply())
    if st == Survey.confirm.state:
        return await message.answer("Проверьте:\n\n" + answers_text(d2), reply_markup=kb_inline("confirm", 2))

    await state.clear()
    await message.answer("Меню:", reply_markup=kb_menu(message.from_user.id))


@router.message(F.text == "📁 Документы")
async def m_docs(message: Message):
    common_path = f"{YANDEX_ROOT}/{YANDEX_LOCAL}"
    try:
        common = await yd.list_files(common_path, limit=30)
    except Exception:
        common = []

    personal = db.list_docs(message.from_user.id, limit=30)

    text = ["📁 Файлы Мастерской '12':"]
    text += [f"- {x['name']}" for x in common] if common else ["- (пусто)"]
    text += ["", "📁 Ваши файлы:"]
    text += [f"- {d['file_name']}" for d in personal] if personal else ["- (пусто)"]

    await message.answer("\n".join(text), reply_markup=kb_menu(message.from_user.id))

    if common:
        kb_rows = []
        for x in common:
            token = os.urandom(6).hex()
            COMMON_DL_MAP[token] = x["path"]
            kb_rows.append([InlineKeyboardButton(text=f"⬇️ {x['name']}", callback_data=f"dlc:{token}")])
        await message.answer("Файлы Мастерской '12' — скачать:", reply_markup=InlineKeyboardMarkup(inline_keyboard=kb_rows))

    if personal:
        kb_rows = [[InlineKeyboardButton(text=f"⬇️ {d['file_name']}", callback_data=f"dlp:{d['id']}")] for d in personal]
        await message.answer("Ваши файлы — скачать:", reply_markup=InlineKeyboardMarkup(inline_keyboard=kb_rows))


@router.callback_query(F.data.startswith("dlc:"))
async def dl_common(call: CallbackQuery):
    token = call.data.split(":", 1)[1]
    path = COMMON_DL_MAP.get(token)
    if not path:
        await call.answer("Кнопка устарела", show_alert=True)
        return
    await call.answer("Скачиваю...")
    try:
        await download_and_send(call.message, path)
    finally:
        COMMON_DL_MAP.pop(token, None)


@router.callback_query(F.data.startswith("dlp:"))
async def dl_personal(call: CallbackQuery):
    doc_id = int(call.data.split(":", 1)[1])

    row = None
    docs = db.list_docs(call.from_user.id, limit=500)
    for d in docs:
        if int(d["id"]) == doc_id:
            row = d
            break

    if row is None and is_admin(call.from_user.id):
        with db._conn() as con:
            row = con.execute("SELECT * FROM docs WHERE id=?", (doc_id,)).fetchone()

    if not row:
        await call.answer("Не найдено", show_alert=True)
        return

    await call.answer("Скачиваю...")
    await download_and_send(call.message, row["ydisk_path"])


@router.callback_query(F.data.startswith("dls:"))
async def dl_submission_file(call: CallbackQuery):
    token = call.data.split(":", 1)[1]
    path = SUB_DL_MAP.get(token)
    if not path:
        await call.answer("Кнопка устарела", show_alert=True)
        return
    await call.answer("Скачиваю...")
    try:
        await download_and_send(call.message, path)
    finally:
        SUB_DL_MAP.pop(token, None)


@router.message(F.text == "📄 Мои ответы")
async def m_my(message: Message):
    last = db.get_last_submission_by_user(message.from_user.id)
    if not last:
        await message.answer("Ответов нет. Нажмите «📝 Пройти опрос».", reply_markup=kb_menu(message.from_user.id))
        return
    a = submission_to_dict(last)
    await message.answer("📄 Последняя анкета:\n\n" + answers_text(a), reply_markup=kb_menu(message.from_user.id))


@router.message(F.text == "✏️ Изменить ответы")
async def m_edit(message: Message, state: FSMContext):
    last = db.get_last_submission_by_user(message.from_user.id)
    if not last:
        await message.answer("Ответов нет. Сначала пройдите опрос.", reply_markup=kb_menu(message.from_user.id))
        return
    await state.clear()
    await state.update_data(edit_sub_id=int(last["id"]))
    await state.set_state(Edit.pick)
    await message.answer("Что изменить?", reply_markup=kb_edit_fields())


@router.message(F.text == "🛠 Админ")
async def m_admin(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    await state.clear()
    await message.answer("Админ-меню:", reply_markup=kb_admin_menu())


@router.message(F.text == "⬅️ Назад")
async def m_back(message: Message, state: FSMContext):
    await state.clear()
    await message.answer("Меню:", reply_markup=kb_menu(message.from_user.id))


@router.message(F.document)
async def on_doc(message: Message, bot: Bot):
    b = BytesIO()
    await bot.download(message.document, destination=b)
    disk_path = await put_to_yandex(
        message.from_user.id,
        message.document.file_name or f"file_{message.document.file_unique_id}",
        b.getvalue(),
    )
    await message.answer(f"✅ Сохранено на Я.Диск:\n{disk_path}", reply_markup=kb_menu(message.from_user.id))


@router.message(F.photo)
async def on_photo(message: Message, bot: Bot):
    ph = message.photo[-1]
    b = BytesIO()
    await bot.download(ph, destination=b)
    disk_path = await put_to_yandex(message.from_user.id, f"photo_{ph.file_unique_id}.jpg", b.getvalue())
    await message.answer(f"✅ Сохранено на Я.Диск:\n{disk_path}", reply_markup=kb_menu(message.from_user.id))


@router.message(F.text == "⏸ Прервать и доделать позже")
async def survey_pause_reply(message: Message, state: FSMContext):
    st = await state.get_state()
    if not st or (not st.startswith("Survey:") and not st.startswith("EditPower:")):
        await message.answer("Опрос сейчас не идет.", reply_markup=kb_menu(message.from_user.id))
        return
    if st.startswith("EditPower:"):
        await message.answer("Сначала завершите изменение «Силовые подключения».", reply_markup=kb_menu(message.from_user.id))
        return
    d = await draft_get(state)
    db.upsert_draft(message.from_user.id, st, json.dumps(d, ensure_ascii=False))
    await state.clear()
    await message.answer("✅ Сохранено. Продолжить: «▶️ Продолжить опрос».", reply_markup=kb_menu(message.from_user.id))


@router.message(Edit.pick)
async def edit_pick(message: Message, state: FSMContext):
    if message.text == "⬅️ Назад":
        await state.clear()
        await message.answer("Меню:", reply_markup=kb_menu(message.from_user.id))
        return

    field = None
    for k, title in EDIT_FIELDS:
        if title == message.text:
            field = k
            break
    if not field:
        await message.answer("Выберите поле кнопкой.")
        return

    if field == "power_block":
        sub_id = (await state.get_data()).get("edit_sub_id")
        row = db.get_submission(int(sub_id)) if sub_id else None
        if not row:
            await state.clear()
            await message.answer("Меню:", reply_markup=kb_menu(message.from_user.id))
            return

        power_where_list = json.loads((row["power_where_json"] or "[]"))
        items_map: Dict[str, List[str]] = {}
        for s in power_where_list:
            if isinstance(s, str) and ": " in s:
                t, w = s.split(": ", 1)
                items_map.setdefault(t, []).append(w)
        power_types = list(items_map.keys())
        power_items = [{"type": t, "count": len(ws), "where": list(ws)} for t, ws in items_map.items()]

        await state.update_data(draft={"power_types": power_types, "power_items": power_items, "power_i": 0})
        await state.set_state(EditPower.power_type)
        await message.answer(
            "Выберите один или несколько вариантов и нажмите «➡️ Далее»:",
            reply_markup=kb_power_types_multi(power_types, none_selected=False),
        )
        return

    await state.update_data(edit_field=field)

    if field in EDIT_OPTIONS:
        await state.set_state(Edit.value)
        await message.answer("Выберите значение:", reply_markup=kb_reply_options(EDIT_OPTIONS[field]))
        return

    await state.set_state(Edit.value)
    hint = "Введите новое значение"
    if field == "event_date":
        hint += " (ДД.ММ.ГГГГ)"
    if field == "phone":
        hint += " (например: +7 999 123-45-67)"
    await message.answer(hint, reply_markup=ReplyKeyboardRemove())


@router.message(Edit.value)
async def edit_value(message: Message, state: FSMContext):
    data = await state.get_data()
    sub_id = data.get("edit_sub_id")
    field = data.get("edit_field")
    if not sub_id or not field:
        await state.clear()
        await message.answer("Меню:", reply_markup=kb_menu(message.from_user.id))
        return

    txt = (message.text or "").strip()

    if txt == "⬅️ Назад":
        await state.set_state(Edit.pick)
        await message.answer("Что изменить?", reply_markup=kb_edit_fields())
        return

    if field in EDIT_OPTIONS and txt not in EDIT_OPTIONS[field]:
        await message.answer("Выберите значение кнопкой.")
        return

    patch: Dict[str, Any] = {}

    if field == "event_date":
        nd = norm_date(txt)
        if not nd:
            await message.answer("Формат даты: ДД.ММ.ГГГГ")
            return
        txt = nd

    if field == "techs_count":
        n = int_pos(txt)
        if n is None:
            await message.answer("Введите число > 0")
            return
        txt = str(n)

    if field == "phone":
        ph = norm_phone(txt)
        if not ph:
            await message.answer("Введите телефон (например: +7 999 123-45-67)")
            return
        txt = ph

    patch[field] = txt

    if field == "extra_equipment" and txt == "Нет":
        patch["plugs"] = "—"
    if field == "dimmer_needed" and txt == "Нет":
        patch["dimmer_text"] = "—"
    if field == "operator" and txt.startswith("Оператор"):
        patch["console_help"] = "—"
        patch["console_model"] = "—"
    if field == "console_help" and txt != "Привезем свой пульт":
        patch["console_model"] = "—"

    if field in {"org", "event_date", "event_title"}:
        row = db.get_submission(int(sub_id))
        if row:
            org = patch.get("org") or row["org"]
            event_date = patch.get("event_date") or row["event_date"]
            event_title = patch.get("event_title") or row["event_title"] or "Мероприятие"
            new_folder = folder_for(event_date, org, event_title)
            await yd.ensure_folder(f"{YANDEX_ROOT}")
            await yd.ensure_folder(new_folder)
            patch["ydisk_folder"] = new_folder
            db.upsert_user_last(message.from_user.id, int(sub_id), new_folder)

    ok = db.update_submission(int(sub_id), patch)
    await state.clear()
    await message.answer("✅ Обновлено" if ok else "Не удалось обновить", reply_markup=kb_menu(message.from_user.id))


@router.message(Survey.org)
async def s_org(message: Message, state: FSMContext):
    await draft_set(state, {"org": (message.text or "").strip()})
    await state.set_state(Survey.role)
    await message.answer("2) Ваша должность?", reply_markup=kb_survey_reply())


@router.message(Survey.role)
async def s_role(message: Message, state: FSMContext):
    await draft_set(state, {"role": (message.text or "").strip()})
    await state.set_state(Survey.name)
    await message.answer("3) Ваше имя?", reply_markup=kb_survey_reply())


@router.message(Survey.name)
async def s_name(message: Message, state: FSMContext):
    await draft_set(state, {"name": (message.text or "").strip()})
    await state.set_state(Survey.event_date)
    await message.answer("4) Дата проведения мероприятия? (ДД.ММ.ГГГГ)", reply_markup=kb_survey_reply())


@router.message(Survey.event_date)
async def s_date(message: Message, state: FSMContext):
    nd = norm_date(message.text)
    if not nd:
        await message.answer("Формат даты: ДД.ММ.ГГГГ", reply_markup=kb_survey_reply())
        return
    await draft_set(state, {"event_date": nd})
    await state.set_state(Survey.event_title)
    await message.answer("5) Название мероприятия?", reply_markup=kb_survey_reply())


@router.message(Survey.event_title)
async def s_event_title(message: Message, state: FSMContext):
    await draft_set(state, {"event_title": (message.text or "").strip()})
    await state.set_state(Survey.scene)
    await message.answer("6) На какой сцене будет проходить мероприятие?", reply_markup=kb_inline("scene", 2))


@router.message(Survey.techs_count)
async def s_techs(message: Message, state: FSMContext):
    n = int_pos(message.text)
    if n is None:
        await message.answer("Введите число > 0", reply_markup=kb_survey_reply())
        return
    await draft_set(state, {"techs_count": str(n)})
    await state.set_state(Survey.extra_equipment)
    await message.answer("10) Используете ли доп. световое оборудование?", reply_markup=kb_inline("extra_equipment", 1))


@router.message(Survey.plugs)
async def s_plugs(message: Message, state: FSMContext):
    await draft_set(state, {"plugs": (message.text or "").strip()})
    await state.set_state(Survey.power_type)
    d = await draft_get(state)
    sel = d.get("power_types") or []
    none_selected = bool(d.get("power_none"))
    await message.answer(
        "12) Выберите один или несколько вариантов и нажмите «➡️ Далее»:",
        reply_markup=kb_power_types_multi(sel, none_selected=none_selected),
    )


@router.message(Survey.power_count)
async def s_power_count(message: Message, state: FSMContext):
    n = int_pos(message.text)
    if n is None:
        await message.answer("Введите число > 0", reply_markup=kb_survey_reply())
        return

    d = await draft_get(state)
    items = d.get("power_items") or []
    i = int(d.get("power_i") or 0)

    if i >= len(items):
        await state.set_state(Survey.dimmer_needed)
        await message.answer("15) Нужен ли диммер?", reply_markup=kb_inline("dimmer_needed", 2))
        return

    items[i]["count"] = int(n)
    items[i]["where"] = []
    await draft_set(state, {"power_items": items})

    await state.set_state(Survey.power_where)
    await message.answer(f"14) Где нужны подключения для «{items[i]['type']}»?", reply_markup=kb_inline("power_where", 2))


@router.message(Survey.dimmer_text)
async def s_dimmer_text(message: Message, state: FSMContext):
    await draft_set(state, {"dimmer_text": (message.text or "").strip()})
    await state.set_state(Survey.sfx)
    sel = (await draft_get(state)).get("sfx_list") or []
    none_selected = bool((await draft_get(state)).get("sfx_none"))
    await message.answer(
        "17) Выберите один или несколько вариантов и нажмите «➡️ Далее»:",
        reply_markup=kb_sfx_multi(sel, none_selected=none_selected),
    )


@router.message(Survey.sfx_other)
async def s_sfx_other(message: Message, state: FSMContext):
    await draft_set(state, {"sfx_other": (message.text or "").strip()})
    await state.set_state(Survey.operator)
    await message.answer("18) Кто ведет мероприятие?", reply_markup=kb_inline("operator", 1))


@router.message(Survey.phone)
async def s_phone(message: Message, state: FSMContext):
    ph = norm_phone(message.text)
    if not ph:
        await message.answer("Введите телефон (например: +7 999 123-45-67)", reply_markup=kb_survey_reply())
        return
    await draft_set(state, {"phone": ph})
    await state.set_state(Survey.confirm)
    d2 = await draft_get(state)
    await message.answer("Проверьте:\n\n" + answers_text(d2), reply_markup=kb_inline("confirm", 2))


@router.message(Survey.console_model)
async def s_console_model(message: Message, state: FSMContext):
    await draft_set(state, {"console_model": (message.text or "").strip()})
    await state.set_state(Survey.phone)
    await message.answer("21) Номер телефона для связи?", reply_markup=kb_survey_reply())


@router.message(EditPower.power_count)
async def ep_power_count(message: Message, state: FSMContext):
    n = int_pos(message.text)
    if n is None:
        await message.answer("Введите число > 0", reply_markup=kb_survey_reply())
        return

    d = await draft_get(state)
    items = d.get("power_items") or []
    i = int(d.get("power_i") or 0)

    if i >= len(items):
        await state.set_state(EditPower.confirm)
        await message.answer("Сохранить изменения силовых?", reply_markup=kb_inline("confirm", 2))
        return

    items[i]["count"] = int(n)
    items[i]["where"] = []
    await draft_set(state, {"power_items": items})

    await state.set_state(EditPower.power_where)
    await message.answer(f"Где нужны подключения для «{items[i]['type']}»?", reply_markup=kb_inline("power_where", 2))


@router.callback_query(F.data.startswith("pt:"))
async def power_types_cb(call: CallbackQuery, state: FSMContext):
    st = await state.get_state()
    if st not in {Survey.power_type.state, EditPower.power_type.state}:
        await call.answer()
        return

    d = await draft_get(state)
    sel: List[str] = list(d.get("power_types") or [])
    opts = SURVEY_OPTIONS["power_type"]

    if call.data == "pt:none":
        sel = []
        await draft_set(state, {"power_types": sel, "power_items": [], "power_i": 0, "power_none": True})
        try:
            await call.message.edit_reply_markup(reply_markup=kb_power_types_multi(sel, none_selected=True))
        except TelegramBadRequest as e:
            if "message is not modified" not in str(e):
                raise
        await call.answer()
        return

    if call.data == "pt:done":
        await draft_set(state, {"power_types": sel})

        if not sel:
            await draft_set(state, {"power_items": [], "power_i": 0})
            if st == EditPower.power_type.state:
                await state.set_state(EditPower.confirm)
                await call.message.answer("Силовые подключения: Нет. Сохранить?", reply_markup=kb_inline("confirm", 2))
            else:
                await state.set_state(Survey.dimmer_needed)
                await call.message.answer("15) Нужен ли диммер?", reply_markup=kb_inline("dimmer_needed", 2))
            await call.answer()
            return

        items = [{"type": t, "count": 0, "where": []} for t in sel]
        await draft_set(state, {"power_items": items, "power_i": 0})
        await state.set_state(EditPower.power_count if st == EditPower.power_type.state else Survey.power_count)
        await call.message.answer(f"13) Сколько подключений нужно для «{items[0]['type']}»?", reply_markup=kb_survey_reply())
        await call.answer()
        return

    m = re.match(r"^pt:opt:(\d+)$", call.data)
    if m:
        idx = int(m.group(1))
        if 0 <= idx < len(opts):
            val = opts[idx]
            if val != "Нет":
                if val in sel:
                    sel.remove(val)
                else:
                    sel.append(val)
        await draft_set(state, {"power_types": sel, "power_none": False})
        try:
            await call.message.edit_reply_markup(reply_markup=kb_power_types_multi(sel, none_selected=False))
        except TelegramBadRequest as e:
            if "message is not modified" not in str(e):
                raise
        await call.answer()
        return

    await call.answer()


@router.callback_query(F.data.startswith("sx:"))
async def sfx_cb(call: CallbackQuery, state: FSMContext):
    st = await state.get_state()
    if st != Survey.sfx.state:
        await call.answer()
        return

    d = await draft_get(state)
    sel: List[str] = list(d.get("sfx_list") or [])
    opts = SURVEY_OPTIONS["sfx"]

    if call.data == "sx:none":
        sel = []
        await draft_set(state, {"sfx_list": sel, "sfx_other": "", "sfx_none": True})
        try:
            await call.message.edit_reply_markup(reply_markup=kb_sfx_multi(sel, none_selected=True))
        except TelegramBadRequest as e:
            if "message is not modified" not in str(e):
                raise
        await call.answer()
        return

    if call.data == "sx:done":
        await draft_set(state, {"sfx_list": sel})
        if not sel:
            await draft_set(state, {"sfx_list": [], "sfx_other": ""})
            await state.set_state(Survey.operator)
            await call.message.answer("18) Кто ведет мероприятие?", reply_markup=kb_inline("operator", 1))
            await call.answer()
            return

        if "Другое" in sel:
            await state.set_state(Survey.sfx_other)
            await call.message.answer("Введите свой вариант спецэффектов:", reply_markup=kb_survey_reply())
            await call.answer()
            return

        await draft_set(state, {"sfx_other": ""})
        await state.set_state(Survey.operator)
        await call.message.answer("18) Кто ведет мероприятие?", reply_markup=kb_inline("operator", 1))
        await call.answer()
        return

    m = re.match(r"^sx:opt:(\d+)$", call.data)
    if m:
        idx = int(m.group(1))
        if 0 <= idx < len(opts):
            val = opts[idx]
            if val != "Нет":
                if val in sel:
                    sel.remove(val)
                else:
                    sel.append(val)
        await draft_set(state, {"sfx_list": sel, "sfx_none": False})
        try:
            await call.message.edit_reply_markup(reply_markup=kb_sfx_multi(sel, none_selected=False))
        except TelegramBadRequest as e:
            if "message is not modified" not in str(e):
                raise
        await call.answer()
        return

    await call.answer()


@router.callback_query(F.data.startswith("ans:"))
async def s_inline(call: CallbackQuery, state: FSMContext):
    _, field, idx_s = call.data.split(":", 2)
    st = await state.get_state()
    d = await draft_get(state)

    if field not in SURVEY_OPTIONS:
        return await call.answer()

    try:
        idx = int(idx_s)
        value = SURVEY_OPTIONS[field][idx]
    except Exception:
        return await call.answer()

    if field == "scene" and st == Survey.scene.state:
        await draft_set(state, {"scene": value})
        await state.set_state(Survey.night_mount)
        await call.message.answer("7) Нужен ли ночной монтаж перед мероприятием?", reply_markup=kb_inline("night_mount", 2))
        return await call.answer()

    if field == "night_mount" and st == Survey.night_mount.state:
        await draft_set(state, {"night_mount": value})
        await state.set_state(Survey.mount_who)
        await call.message.answer("8) Кто производит монтаж световой аппаратуры?", reply_markup=kb_inline("mount_who", 1))
        return await call.answer()

    if field == "mount_who" and st == Survey.mount_who.state:
        await draft_set(state, {"mount_who": value})
        await state.set_state(Survey.techs_count)
        await call.message.answer("9) Сколько техников на монтаж понадобится?", reply_markup=kb_survey_reply())
        return await call.answer()

    if field == "extra_equipment" and st == Survey.extra_equipment.state:
        await draft_set(state, {"extra_equipment": value})
        if value == "Нет":
            await draft_set(state, {"plugs": "—"})
            await state.set_state(Survey.power_type)
            sel = (await draft_get(state)).get("power_types") or []
            none_selected = bool((await draft_get(state)).get("power_none"))
            await call.message.answer(
                "12) Выберите один или несколько вариантов и нажмите «➡️ Далее»:",
                reply_markup=kb_power_types_multi(sel, none_selected=none_selected),
            )
        else:
            await state.set_state(Survey.plugs)
            await call.message.answer("11) Какие вилки на ваших приборах?", reply_markup=kb_survey_reply())
        return await call.answer()

    if field == "power_where" and st in {Survey.power_where.state, EditPower.power_where.state}:
        items = d.get("power_items") or []
        i = int(d.get("power_i") or 0)

        if i >= len(items):
            if st == EditPower.power_where.state:
                await state.set_state(EditPower.confirm)
                await call.message.answer("Сохранить изменения силовых?", reply_markup=kb_inline("confirm", 2))
            else:
                await state.set_state(Survey.dimmer_needed)
                await call.message.answer("15) Нужен ли диммер?", reply_markup=kb_inline("dimmer_needed", 2))
            return await call.answer()

        cur = items[i]
        need = int(cur.get("count") or 0)
        wh = cur.get("where") or []

        if len(wh) < need:
            wh.append(value)
            cur["where"] = wh
            items[i] = cur
            await draft_set(state, {"power_items": items})

        if len(wh) < need:
            left = need - len(wh)
            await call.message.answer(
                f"14) Где нужны подключения для «{cur['type']}»? Осталось {left}",
                reply_markup=kb_inline("power_where", 2),
            )
            return await call.answer()

        i += 1
        if i < len(items):
            await draft_set(state, {"power_i": i})
            await state.set_state(EditPower.power_count if st == EditPower.power_where.state else Survey.power_count)
            await call.message.answer(f"13) Сколько подключений нужно для «{items[i]['type']}»?", reply_markup=kb_survey_reply())
            return await call.answer()

        if st == EditPower.power_where.state:
            await state.set_state(EditPower.confirm)
            parts = []
            for it in items:
                parts.append(f"{it['type']} — {it.get('count', 0)}: " + ", ".join(it.get("where") or []))
            await call.message.answer("Силовые подключения:\n" + "\n".join(parts) + "\n\nСохранить?", reply_markup=kb_inline("confirm", 2))
        else:
            await state.set_state(Survey.dimmer_needed)
            await call.message.answer("15) Нужен ли диммер?", reply_markup=kb_inline("dimmer_needed", 2))
        return await call.answer()

    if field == "dimmer_needed" and st == Survey.dimmer_needed.state:
        await draft_set(state, {"dimmer_needed": value})
        if value == "Нет":
            await draft_set(state, {"dimmer_text": "—"})
            await state.set_state(Survey.sfx)
            sel = (await draft_get(state)).get("sfx_list") or []
            none_selected = bool((await draft_get(state)).get("sfx_none"))
            await call.message.answer(
                "17) Выберите один или несколько вариантов и нажмите «➡️ Далее»:",
                reply_markup=kb_sfx_multi(sel, none_selected=none_selected),
            )
        else:
            await state.set_state(Survey.dimmer_text)
            await call.message.answer("16) Диммер: где и сколько?", reply_markup=kb_survey_reply())
        return await call.answer()

    if field == "operator" and st == Survey.operator.state:
        await draft_set(state, {"operator": value})
        if value.startswith("Оператор"):
            await draft_set(state, {"console_help": "—", "console_model": "—"})
            await state.set_state(Survey.phone)
            await call.message.answer("21) Номер телефона для связи?", reply_markup=kb_survey_reply())
        else:
            await state.set_state(Survey.console_help)
            d2 = await draft_get(state)
            console_name = "GrandMa2 Light" if d2.get("scene") == "Большой зал" else "Chamsys MQ500"
            await call.message.answer(
                f"19) Мы используем пульт {console_name}. Нужна помощь с пультом?",
                reply_markup=kb_inline("console_help", 1),
            )
        return await call.answer()

    if field == "console_help" and st == Survey.console_help.state:
        await draft_set(state, {"console_help": value})
        if value == "Привезем свой пульт":
            await state.set_state(Survey.console_model)
            await call.message.answer("20) Марка и модель вашего пульта? (текст)", reply_markup=kb_survey_reply())
        else:
            await draft_set(state, {"console_model": "—"})
            await state.set_state(Survey.phone)
            await call.message.answer("21) Номер телефона для связи?", reply_markup=kb_survey_reply())
        return await call.answer()

    if field == "confirm" and st in {Survey.confirm.state, EditPower.confirm.state}:
        if st == EditPower.confirm.state:
            if value.startswith("🔁"):
                await state.clear()
                await call.message.answer("Меню:", reply_markup=kb_menu(call.from_user.id))
                return await call.answer()

            d2 = await draft_get(state)
            power_items = d2.get("power_items") or []
            power_types = d2.get("power_types") or []

            if not power_items and power_types:
                power_items = [{"type": t, "count": 0, "where": []} for t in power_types]

            if not power_items:
                power_type = "Нет"
                power_where_list = []
                power_count_sum = 0
            else:
                power_type = ", ".join([x.get("type") for x in power_items if x.get("type")])
                power_where_list = []
                power_count_sum = 0
                for it in power_items:
                    power_count_sum += int(it.get("count") or 0)
                    t = it.get("type")
                    for w in (it.get("where") or []):
                        power_where_list.append(f"{t}: {w}")

            power_where_json = json.dumps(power_where_list, ensure_ascii=False)

            data = await state.get_data()
            sub_id = data.get("edit_sub_id")
            if not sub_id:
                await state.clear()
                await call.message.answer("Меню:", reply_markup=kb_menu(call.from_user.id))
                return await call.answer()

            db.update_submission(
                int(sub_id),
                {"power_type": power_type, "power_count": str(power_count_sum), "power_where_json": power_where_json},
            )

            await state.clear()
            await call.message.answer("✅ Силовые обновлены", reply_markup=kb_menu(call.from_user.id))
            return await call.answer()

        if value.startswith("🔁"):
            db.delete_draft(call.from_user.id)
            await state.clear()
            await call.message.answer("Меню:", reply_markup=kb_menu(call.from_user.id))
            return await call.answer()

        d2 = await draft_get(state)
        d2.setdefault("plugs", "—")
        d2.setdefault("dimmer_needed", "Нет")
        d2.setdefault("dimmer_text", "—")
        d2.setdefault("console_help", "—")
        d2.setdefault("console_model", "—")
        d2.setdefault("event_title", "Мероприятие")
        d2.setdefault("sfx_list", [])
        d2.setdefault("sfx_other", "")
        d2.setdefault("phone", "")

        power_items = d2.get("power_items") or []
        power_types = d2.get("power_types") or []
        if not power_items and power_types:
            power_items = [{"type": t, "count": 0, "where": []} for t in power_types]

        if not power_items:
            d2["power_type"] = "Нет"
            power_where_list = []
            power_count_sum = 0
        else:
            d2["power_type"] = ", ".join([x.get("type") for x in power_items if x.get("type")])
            power_where_list = []
            power_count_sum = 0
            for it in power_items:
                power_count_sum += int(it.get("count") or 0)
                for w in (it.get("where") or []):
                    power_where_list.append(f"{it.get('type')}: {w}")

        sfx_list = d2.get("sfx_list") or []
        d2["sfx_json"] = json.dumps(sfx_list, ensure_ascii=False)
        d2["sfx_other"] = str(d2.get("sfx_other") or "").strip()

        folder = folder_for(d2["event_date"], d2["org"], d2["event_title"])
        await yd.ensure_folder(f"{YANDEX_ROOT}")
        await yd.ensure_folder(folder)

        payload = dict(d2)
        payload["power_count"] = str(power_count_sum)
        payload["power_where_json"] = json.dumps(power_where_list, ensure_ascii=False)
        payload["ydisk_folder"] = folder
        payload["sfx_json"] = d2["sfx_json"]
        payload["sfx_other"] = d2["sfx_other"]

        sub_id = db.insert_submission(call.from_user.id, payload)
        db.upsert_user_last(call.from_user.id, sub_id, folder)
        db.delete_draft(call.from_user.id)

        # === АВТОГЕНЕРАЦИЯ WORD + ЗАГРУЗКА В ПАПКУ АНКЕТЫ ===
        try:
            sub_for_doc = dict(payload)
            sub_for_doc["id"] = sub_id
            sub_for_doc["user_id"] = call.from_user.id
            sub_for_doc["ydisk_folder"] = folder

            # нормализуем console_model как в админ-выгрузке
            if sub_for_doc.get("operator") != "Оператор Мастерской «12»":
                if sub_for_doc.get("console_help") == "Привезем свой пульт":
                    sub_for_doc["console_model"] = sub_for_doc.get("console_model") or "—"
                else:
                    sub_for_doc["console_model"] = (
                        "GrandMa2 Light" if sub_for_doc.get("scene") == "Большой зал" else "Chamsys MQ500"
                    )
            else:
                sub_for_doc["console_model"] = "—"

            docx_data = build_docx_for_submission(sub_for_doc)
            docx_name = (
                f"анкета_{sub_id}_"
                f"{sanitize_name(sub_for_doc['event_date'])}_"
                f"{sanitize_name(sub_for_doc['org'])}_"
                f"{sanitize_name(sub_for_doc['event_title'])}.docx"
            )
            await put_to_yandex(call.from_user.id, docx_name, docx_data)
        except Exception:
            # не валим сохранение анкеты из-за docx
            pass
        # === /АВТОГЕНЕРАЦИЯ WORD ===

        await state.clear()
        await call.message.answer(THANKS, reply_markup=kb_menu(call.from_user.id))
        await call.message.answer("✅ Анкета сохранена", reply_markup=kb_menu(call.from_user.id))
        return await call.answer()

    await call.answer()


@router.message(
    F.state.in_(
        {
            Survey.scene,
            Survey.night_mount,
            Survey.mount_who,
            Survey.extra_equipment,
            Survey.dimmer_needed,
            Survey.operator,
            Survey.console_help,
            Survey.sfx,
        }
    )
)
async def inline_only_text(message: Message):
    await message.answer("Выберите вариант кнопками выше.")


@router.message(F.state.in_({Survey.power_type, EditPower.power_type}))
async def inline_power_type_text(message: Message):
    await message.answer("Выберите варианты кнопками выше.")


def month_shift(year: int, month: int, delta: int) -> Tuple[int, int]:
    m = month + delta
    y = year
    while m <= 0:
        m += 12
        y -= 1
    while m > 12:
        m -= 12
        y += 1
    return y, m


def kb_months() -> ReplyKeyboardMarkup:
    now = datetime.now()
    y0, m0 = now.year, now.month
    y_prev, m_prev = month_shift(y0, m0, -1)
    y_next, m_next = month_shift(y0, m0, +1)

    def t(y: int, m: int) -> str:
        return f"{y}-{m:02d}"

    return ReplyKeyboardMarkup(
        keyboard=[
            [
                KeyboardButton(text=f"◀️ {t(y_prev, m_prev)}"),
                KeyboardButton(text=f"✅ {t(y0, m0)}"),
                KeyboardButton(text=f"▶️ {t(y_next, m_next)}"),
            ],
            [KeyboardButton(text="⬅️ Назад")],
        ],
        resize_keyboard=True,
    )


def parse_month_btn(s: str) -> Optional[Tuple[int, int]]:
    s = (s or "").strip()
    m = re.search(r"(\d{4})-(\d{2})", s)
    if not m:
        return None
    y = int(m.group(1))
    mm = int(m.group(2))
    if mm < 1 or mm > 12:
        return None
    return y, mm


def kb_forms_list(rows: List[Any]) -> ReplyKeyboardMarkup:
    kb_rows: List[List[KeyboardButton]] = []
    for r in rows:
        title = f"#{r['id']} {r['event_date']} | {r['org']} | {r['event_title']}"
        kb_rows.append([KeyboardButton(text=title[:120])])
    kb_rows.append([KeyboardButton(text="⬅️ Назад")])
    return ReplyKeyboardMarkup(keyboard=kb_rows, resize_keyboard=True)


@router.message(F.text == "📄 Word анкеты")
async def a_word_start(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    await state.clear()
    await state.set_state(AdminWord.pick_month)
    await message.answer("Выберите месяц:", reply_markup=kb_months())


@router.message(AdminWord.pick_month)
async def a_word_pick_month(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    if message.text == "⬅️ Назад":
        await state.clear()
        await message.answer("Админ-меню:", reply_markup=kb_admin_menu())
        return

    ym = parse_month_btn(message.text or "")
    if not ym:
        await message.answer("Выберите месяц кнопкой.", reply_markup=kb_months())
        return

    year, month = ym
    rows = db.list_submissions_by_month(year, month, limit=200)
    if not rows:
        await message.answer("Анкет за этот месяц нет.", reply_markup=kb_months())
        return

    await state.set_state(AdminWord.pick_form)
    await message.answer("Выберите анкету:", reply_markup=kb_forms_list(rows))


@router.message(AdminWord.pick_form)
async def a_word_pick_form(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    if message.text == "⬅️ Назад":
        await state.set_state(AdminWord.pick_month)
        await message.answer("Выберите месяц:", reply_markup=kb_months())
        return

    m = re.match(r"^#(\d+)", (message.text or "").strip())
    if not m:
        await message.answer("Выберите анкету кнопкой.")
        return

    sub_id = int(m.group(1))
    sub_row = db.get_submission(sub_id)
    if not sub_row:
        await message.answer("Не найдено.")
        return

    sub = dict(sub_row)

    if sub.get("operator") != "Оператор Мастерской «12»":
        if sub.get("console_help") == "Привезем свой пульт":
            sub["console_model"] = sub.get("console_model") or "—"
        else:
            sub["console_model"] = "GrandMa2 Light" if sub.get("scene") == "Большой зал" else "Chamsys MQ500"
    else:
        sub["console_model"] = "—"

    data = build_docx_for_submission(sub)
    filename = f"анкета_{sub_id}_{sanitize_name(sub['event_date'])}_{sanitize_name(sub['org'])}_{sanitize_name(sub['event_title'])}.docx"
    await message.answer_document(
        BufferedInputFile(data, filename=filename),
        caption=f"Анкета #{sub_id}",
        reply_markup=kb_admin_menu(),
    )
    await state.clear()


@router.message(F.text == "🗑 Удалить анкету")
async def a_del_start(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    await state.clear()
    await state.set_state(AdminDel.pick_month)
    await message.answer("Выберите месяц для удаления:", reply_markup=kb_months())


@router.message(AdminDel.pick_month)
async def a_del_pick_month(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    if message.text == "⬅️ Назад":
        await state.clear()
        await message.answer("Админ-меню:", reply_markup=kb_admin_menu())
        return

    ym = parse_month_btn(message.text or "")
    if not ym:
        await message.answer("Выберите месяц кнопкой.", reply_markup=kb_months())
        return

    year, month = ym
    rows = db.list_submissions_by_month(year, month, limit=200)
    if not rows:
        await message.answer("Анкет за этот месяц нет.", reply_markup=kb_months())
        return

    await state.set_state(AdminDel.pick_form)
    await message.answer("Выберите анкету для удаления:", reply_markup=kb_forms_list(rows))


@router.message(AdminDel.pick_form)
async def a_del_pick_form(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    if message.text == "⬅️ Назад":
        await state.set_state(AdminDel.pick_month)
        await message.answer("Выберите месяц:", reply_markup=kb_months())
        return

    m = re.match(r"^#(\d+)", (message.text or "").strip())
    if not m:
        await message.answer("Выберите анкету кнопкой.")
        return

    sub_id = int(m.group(1))
    sub = db.get_submission(sub_id)
    if not sub:
        await message.answer("Не найдено.")
        return

    await state.update_data(del_sub_id=sub_id)
    await state.set_state(AdminDel.confirm)

    kb = ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text="✅ Удалить"), KeyboardButton(text="⬅️ Назад")]],
        resize_keyboard=True,
    )
    await message.answer(
        "Точно удалить анкету?\n\n"
        f"#{sub_id} | {sub['event_date']} | {sub['org']} | {sub['event_title']}\n"
        f"Папка: {sub['ydisk_folder']}",
        reply_markup=kb,
    )


@router.message(AdminDel.confirm)
async def a_del_confirm(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return

    if message.text == "⬅️ Назад":
        await state.set_state(AdminDel.pick_month)
        await message.answer("Выберите месяц:", reply_markup=kb_months())
        return

    if message.text != "✅ Удалить":
        await message.answer("Нажмите «✅ Удалить» или «⬅️ Назад».")
        return

    data = await state.get_data()
    sub_id = int(data.get("del_sub_id") or 0)
    if not sub_id:
        await state.clear()
        await message.answer("Админ-меню:", reply_markup=kb_admin_menu())
        return

    sub = db.get_submission(sub_id)
    folder = sub["ydisk_folder"] if sub else None

    if folder:
        try:
            await yd.delete(folder, permanently=False)
        except Exception:
            pass

    ok = db.delete_submission(sub_id)

    await state.clear()
    await message.answer("✅ удалено" if ok else "не найдено", reply_markup=kb_admin_menu())


@router.message(F.text == "📊 Статистика")
async def a_stats(message: Message):
    if not is_admin(message.from_user.id):
        return
    await message.answer(f"Всего анкет: {db.count_submissions()}", reply_markup=kb_admin_menu())


@router.message(F.text == "📋 Анкеты")
async def a_forms_start(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return
    await state.clear()
    await state.set_state(AdminForms.pick_month)
    await message.answer("Выберите месяц:", reply_markup=kb_months())


@router.message(AdminForms.pick_month)
async def a_forms_pick_month(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return

    if message.text == "⬅️ Назад":
        await state.clear()
        await message.answer("Админ-меню:", reply_markup=kb_admin_menu())
        return

    ym = parse_month_btn(message.text or "")
    if not ym:
        await message.answer("Выберите месяц кнопкой.", reply_markup=kb_months())
        return

    year, month = ym
    rows = db.list_submissions_by_month(year, month, limit=200)
    if not rows:
        await message.answer("Анкет за этот месяц нет.", reply_markup=kb_months())
        return

    await state.set_state(AdminForms.pick_form)
    await message.answer("Выберите анкету:", reply_markup=kb_forms_list(rows))


@router.message(AdminForms.pick_form)
async def a_forms_pick_form(message: Message, state: FSMContext):
    if not is_admin(message.from_user.id):
        return

    if message.text == "⬅️ Назад":
        await state.set_state(AdminForms.pick_month)
        await message.answer("Выберите месяц:", reply_markup=kb_months())
        return

    m = re.match(r"^#(\d+)", (message.text or "").strip())
    if not m:
        await message.answer("Выберите анкету кнопкой.")
        return

    sub_id = int(m.group(1))
    sub = db.get_submission(sub_id)
    if not sub:
        await message.answer("Не найдено.")
        return

    a = submission_to_dict(sub)
    info = "📄 Анкета\n\n" + answers_text(a) + f"\n\n📂 {sub['ydisk_folder']}"
    await message.answer(info, reply_markup=kb_admin_menu())

    try:
        files = await yd.list_files(sub["ydisk_folder"], limit=50)
    except Exception:
        files = []

    if not files:
        await message.answer("Файлов в папке нет.")
        return

    kb_rows = []
    for f in files:
        token = os.urandom(6).hex()
        SUB_DL_MAP[token] = f["path"]
        kb_rows.append([InlineKeyboardButton(text=f"⬇️ {f['name']}", callback_data=f"dls:{token}")])

    await message.answer("Документы анкеты — скачать:", reply_markup=InlineKeyboardMarkup(inline_keyboard=kb_rows))


@router.message()
async def fallback(message: Message):
    await message.answer("Меню:", reply_markup=kb_menu(message.from_user.id))


async def main():
    bot = Bot(token=BOT_TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))
    await yd.ensure_folder(f"{YANDEX_ROOT}")
    await yd.ensure_folder(f"{YANDEX_ROOT}/{YANDEX_LOCAL}")
    await yd.ensure_folder(f"{YANDEX_ROOT}/{YANDEX_INBOX}")
    dp = Dispatcher()
    dp.include_router(router)
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())